from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "training-package"


def set_cell(table, row, col, text):
    table.cell(row, col).text = text


def update_uat():
    path = OUT / "ATEC-Production-Role-UAT-Signoff-Checklist.docx"
    doc = Document(path)
    control = doc.tables[0]
    set_cell(control, 2, 1, "1.0")
    set_cell(control, 5, 1, "10 August 2026")
    set_cell(control, 6, 1, "Codex-assisted production verification")
    set_cell(control, 7, 1, "ATEC Process Owner - signature pending")
    doc.sections[0].header.paragraphs[0].text = "ATEC-UAT-TIME-002 | Production Role-UAT Sign-off Checklist | v1.0"
    doc.paragraphs[3].text = "☒ Process owner approved the UAT scenario and synthetic replacement job."
    doc.paragraphs[4].text = "☒ Authorised Assistant, Inspector and Manager training accounts were available."
    doc.paragraphs[5].text = "☒ ATEC Training customer/site was used without altering operational records."
    doc.paragraphs[6].text = "☐ Desktop browser passed; approved physical mobile device remains open."
    doc.paragraphs[7].text = "☒ The tester used the detailed role guide and Tutorials 9-12."

    public = doc.tables[2]
    for row in range(1, 4):
        set_cell(public, row, 2, "PASS")
    set_cell(public, 1, 3, "HTTPS login loaded successfully.")
    set_cell(public, 2, 3, "Backend responded: ATEC backend is running.")
    set_cell(public, 3, 3, "Refresh and authenticated navigation passed.")

    assistant = doc.tables[3]
    for row in (1, 2, 3, 7):
        set_cell(assistant, row, 2, "PASS")
    for row in (4, 5, 6):
        set_cell(assistant, row, 2, "NOT RUN")
    set_cell(assistant, 1, 3, "01Assistant identity and ASSISTANT role verified.")
    set_cell(assistant, 2, 3, "Only My Day / Timesheet and Account exposed.")
    set_cell(assistant, 3, 3, "Copied-time, exception, history, PDF and submit controls verified.")
    set_cell(assistant, 4, 3, "Manual exception save left for supervised competency exercise.")
    set_cell(assistant, 5, 3, "Negative validation test left for supervised competency exercise.")
    set_cell(assistant, 6, 3, "Correction audit test left for supervised competency exercise.")
    set_cell(assistant, 7, 3, "2.00-hour synthetic sheet submitted; controls became read-only.")

    inspector = doc.tables[4]
    for row in range(1, 8):
        set_cell(inspector, row, 2, "PASS")
    inspector_evidence = {
        1: "01Inspector identity and INSPECTOR role verified.",
        2: "Job Cards and My Day available; approvals absent.",
        3: "Created controlled JC-2026-00004 under ATEC Training.",
        4: "01Assistant included once with Assistant role.",
        5: "Crew persisted after save/reopen.",
        6: "08:00-10:00 factual synthetic sequence retained.",
        7: "Three segments copied to included Assistant on submission.",
    }
    for row, text in inspector_evidence.items():
        set_cell(inspector, row, 3, text)

    copied = doc.tables[5]
    for row in (1, 3, 4):
        set_cell(copied, row, 2, "PASS")
    set_cell(copied, 2, 2, "NOT RUN")
    set_cell(copied, 1, 3, "TRAVEL 0.50 + WORK 1.00 + TRAVEL 0.50 matched job 9990810.")
    set_cell(copied, 2, 3, "No correction was required; retained for competency exercise.")
    set_cell(copied, 3, 3, "Exactly three unique copied rows; no duplicate segment.")
    set_cell(copied, 4, 3, "Status changed to EMPLOYEE SUBMITTED and became read-only.")

    manager = doc.tables[6]
    for row in (1, 2, 4, 5):
        set_cell(manager, row, 2, "PASS")
    set_cell(manager, 3, 2, "NOT RUN")
    set_cell(manager, 1, 3, "01Manager identity verified; approval queue available.")
    set_cell(manager, 2, 3, "01Assistant, 10 Aug 2026, job 9990810 and totals verified.")
    set_cell(manager, 3, 3, "Return path retained for supervised Manager competency exercise.")
    set_cell(manager, 4, 3, "Status changed to MANAGER APPROVED.")
    set_cell(manager, 5, 3, "Inspector menu contained no Timesheet Approvals function.")

    mobile = doc.tables[7]
    for row in range(1, 4):
        set_cell(mobile, row, 3, "OPEN")
        set_cell(mobile, row, 4, "Physical approved-device check remains a rollout hold point.")

    for index, text in {
        15: "☒ Cancel synthetic job JC-2026-00004 after the controlled approval-path test.",
        16: "☒ No temporary accounts were created; controlled role accounts retained.",
        17: "☒ Confirm no payroll export, HR acceptance or Accelo package was created.",
        18: "☐ Store any process-owner-approved production screenshots in the controlled evidence location.",
        19: "☒ Record defects and completed fixes in the improvement register.",
    }.items():
        doc.paragraphs[index].text = text

    decision = doc.tables[8]
    set_cell(decision, 1, 1, "☐")
    set_cell(decision, 2, 1, "☒")
    set_cell(decision, 3, 1, "☐")
    doc.paragraphs[21].text = (
        "Open actions / conditions:\n"
        "1. Process-owner and document-control signatures.\n"
        "2. Physical approved-device checks for Login, Assistant My Day and Inspector Job Crew.\n"
        "3. Manual exception, negative validation, correction-audit and Manager-return exercises during competency sign-off."
    )
    doc.save(path)


def update_manifest():
    path = OUT / "ATEC-Training-Final-Release-Manifest-Improvement-Register.docx"
    doc = Document(path)
    control = doc.tables[0]
    set_cell(control, 2, 1, "1.0 CONTROLLED RELEASE")
    set_cell(control, 5, 1, "Desktop production workflow approved with rollout conditions; physical mobile checks and signatures remain open")
    doc.sections[0].header.paragraphs[0].text = "ATEC-TRN-MAN-001 | Final Release Manifest & Improvement Register | v1.0"
    doc.sections[0].footer.paragraphs[0].text = "CONTROLLED DESKTOP TRAINING RELEASE - COMPLETE ROLLOUT HOLD POINTS BEFORE FINAL SIGN-OFF"

    boundary = doc.tables[1]
    set_cell(boundary, 0, 0, (
        "VERIFICATION BOUNDARY: Repository checks, database UAT, deployment, authenticated desktop role UAT and a controlled end-to-end crew/time/approval transaction are complete. "
        "JC-2026-00004 was cancelled after testing; its timesheet remains Manager Approved as audit evidence and was not HR-accepted or exported. "
        "Physical mobile-device checks and controlled signatures remain rollout hold points."
    ))

    evidence = doc.tables[4]
    set_cell(evidence, 4, 1, "PASS - deployed production build")
    set_cell(evidence, 4, 2, "Commit 00c047e9 live verified")
    set_cell(evidence, 5, 1, "PASS - production 01Assistant")
    set_cell(evidence, 5, 2, "Restricted menu, copied time, submission and read-only state")
    set_cell(evidence, 6, 1, "PASS - production 01Inspector")
    set_cell(evidence, 6, 2, "Job creation, crew persistence, timeline and submission")
    set_cell(evidence, 7, 1, "PASS - authenticated and public")
    set_cell(evidence, 7, 2, "Login plus Assistant, Inspector, Manager and Admin sessions")
    set_cell(evidence, 8, 1, "PASS - ATEC backend is running")
    set_cell(evidence, 8, 2, "Public health and authenticated API-backed screens")
    set_cell(evidence, 11, 1, "Rebuilt after controlled-release document update")
    set_cell(evidence, 11, 2, "ATEC-Complete-Training-Package-v1.0-CONTROLLED-RELEASE.zip")

    improvements = doc.tables[5]
    set_cell(improvements, 1, 3, "DEPLOYED - crew locked after submission; warning live verified (6658255f).")
    set_cell(improvements, 2, 3, "COMPLETE - authenticated roles and controlled JC-2026-00004 workflow passed.")
    set_cell(improvements, 3, 3, "OPEN ROLLOUT HOLD POINT - approved physical mobile device required.")
    set_cell(improvements, 4, 3, "OPEN - process-owner-approved production screenshots remain optional controlled evidence.")
    set_cell(improvements, 8, 3, "OPEN - replace draft images only after screenshot approval.")
    improvement_rows = {row.cells[0].text.strip(): index for index, row in enumerate(improvements.rows)}
    for item in (
        ["TIME-004", "P0", "Realign open timesheets when an employee's approving Manager changes.", "DEPLOYED and production verified (00c047e9)."],
        ["DATE-001", "P0", "Return approval/history dates as database date text to prevent one-day UTC shifts.", "DEPLOYED and production verified (00c047e9)."],
    ):
        if item[0] not in improvement_rows:
            improvements.add_row()
            improvement_rows[item[0]] = len(improvements.rows) - 1
        row_index = improvement_rows[item[0]]
        for col, value in enumerate(item):
            set_cell(improvements, row_index, col, value)

    conditions = doc.tables[6]
    set_cell(conditions, 5, 1, "Satisfied - authenticated role and transaction UAT passed")
    set_cell(conditions, 6, 1, "Open rollout hold point")
    set_cell(conditions, 7, 1, "Open controlled signatures")
    set_cell(conditions, 8, 1, "Satisfied - commits 6658255f and 00c047e9 deployed and verified")
    set_cell(doc.tables[7], 0, 0, (
        "RELEASE DECISION: Accepted as a CONTROLLED DESKTOP TRAINING RELEASE with minor rollout actions. "
        "Use the approved workflow and role guides for desktop training. Complete physical mobile checks, supervised negative/correction exercises and controlled signatures before organisation-wide final sign-off."
    ))
    doc.save(path)


if __name__ == "__main__":
    update_uat()
    update_manifest()
    print("Updated controlled UAT checklist and final release manifest.")
